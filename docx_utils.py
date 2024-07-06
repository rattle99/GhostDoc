import os
import re
from io import BytesIO
from xml.etree.ElementTree import XML
from zipfile import ZipFile

import requests
from bs4 import BeautifulSoup
from bs4 import NavigableString
from bs4 import Tag
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from tika import parser
from werkzeug.utils import secure_filename


def get_html_from_docx(file_path):
    """

    :param file_path: 

    """
    # Parse the DOCX file
    parsed = parser.from_file(file_path, xmlContent=True)

    # Extract HTML content
    html_content = parsed.get("content")

    return html_content


def extractText(html_content):
    """

    :param html_content: 

    """
    soup = BeautifulSoup(html_content, "html.parser")

    def getText(element):
        """

        :param element: 

        """
        if isinstance(element, NavigableString):
            return str(element)
        elif isinstance(element, Tag):
            text = ""
            for child in element.contents:
                text += getText(child)
            return text

    return getText(soup)


def replace_substrings(main_string, replacements):
    """

    :param main_string: 
    :param replacements: 

    """
    # pattern = re.compile("|".join(re.escape(key) for key in replacements.keys()))
    pattern = re.compile(
        r"\b(" + "|".join(re.escape(key) for key in replacements.keys()) + r")\b"
    )

    # Function to replace matched substrings using the dictionary
    def replace_match(match):
        """

        :param match: 

        """
        return replacements[match.group(0)]

    # Use sub method to replace all matches in one go
    result = pattern.sub(replace_match, main_string)
    return result


def modify_html_content(mapDict, html_content):
    """

    :param mapDict: 
    :param html_content: 

    """
    print(len(mapDict))
    soup = BeautifulSoup(html_content, "html.parser")

    def iterate_html(element):
        """

        :param element: 

        """
        if isinstance(element, NavigableString):
            # print(element.string)
            # print(replace_substrings(element.string, mapDict))
            p = False
            if "Rukh Khan" in str(element.string):
                print(element.string)
                p = True
            new_string = replace_substrings(element.string, mapDict)
            element.replace_with(new_string)
            if p:
                print(element.string)
                print(mapDict["Rukh Khan"])
                print("\n\n")
            # print(element.string)
            # print("\n\n")

        elif isinstance(element, Tag):
            for child in element.contents:
                iterate_html(child)

    iterate_html(soup)
    return str(soup)


def export_to_docx(original_file_path, mapDict, uploadFolder):
    """

    :param original_file_path: 
    :param mapDict: 
    :param uploadFolder: 

    """
    original_ext = original_file_path.rsplit(".", 1)[-1].lower()
    temp_file = secure_filename("modified_" + os.path.basename(original_file_path))
    temp_file_path = os.path.join(uploadFolder, temp_file)

    html_content = get_html_from_docx(original_file_path)
    modified_html_content = modify_html_content(mapDict, html_content)
    html_to_docx(original_file_path, modified_html_content, temp_file_path)
    return temp_file_path, "application/msword"


def extract_embedded_images(docx_file):
    """

    :param docx_file: 

    """
    embedded_images = {}

    with ZipFile(docx_file, "r") as docx:
        for entry in docx.namelist():
            if entry.startswith("word/media/"):
                image_data = docx.read(entry)
                embedded_images[entry] = image_data

    return embedded_images


def add_hyperlink(paragraph, url, text):
    """

    :param paragraph: 
    :param url: 
    :param text: 

    """
    # This function adds a hyperlink to a paragraph.
    part = paragraph.part
    r_id = part.relate_to(
        url,
        qn(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
        ),
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(
        qn("r:id"),
        r_id,
    )

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # This makes the text appear blue and underlined
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._element.append(hyperlink)


def add_html_to_docx(html_content, doc, embedded_images):
    """

    :param html_content: 
    :param doc: 
    :param embedded_images: 

    """
    soup = BeautifulSoup(html_content, "html.parser")

    def add_element_to_paragraph(element, paragraph):
        """

        :param element: 
        :param paragraph: 

        """
        if isinstance(element, str):
            paragraph.add_run(element)
        elif element.name == "b":
            run = paragraph.add_run(element.get_text())
            run.bold = True
        elif element.name == "i":
            run = paragraph.add_run(element.get_text())
            run.italic = True
        elif element.name == "u":
            run = paragraph.add_run(element.get_text())
            run.underline = True
        elif element.name == "a":
            run = paragraph.add_run(element.get_text())
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.underline = True
            # add_hyperlink(paragraph, element['href'], element.get_text())
        elif element.name == "img":
            try:
                img_url = element["src"]
                if img_url.startswith("data:image"):
                    image_data = re.sub("^data:image/.+;base64,", "", img_url)
                    image = BytesIO(base64.b64decode(image_data))
                    doc.add_picture(image)
                elif img_url.startswith("embedded:"):
                    image_key = "word/media/" + img_url.split(":")[1]
                    if image_key in embedded_images:
                        image = BytesIO(embedded_images[image_key])
                        doc.add_picture(image)
                else:
                    response = requests.get(img_url)
                    image = BytesIO(response.content)
                    doc.add_picture(image)
            except Exception as e:
                print(f"Error adding image ({img_url}): {e}")

    def add_elements_to_doc(elements, doc):
        """

        :param elements: 
        :param doc: 

        """
        for element in elements:
            if element.name == "p":
                paragraph = doc.add_paragraph()
                for child in element.children:
                    add_element_to_paragraph(child, paragraph)
            elif element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                level = int(element.name[1])
                doc.add_heading(element.get_text(), level=level)
            elif element.name == "img":
                add_element_to_paragraph(element, doc.add_paragraph())
            elif isinstance(element, str):
                doc.add_paragraph(element)

    add_elements_to_doc(soup.body.contents, doc)


def html_to_docx(original_file_path, html_content, output_file):
    """

    :param original_file_path: 
    :param html_content: 
    :param output_file: 

    """
    # Create a new Document
    doc = Document()

    # Extract embedded images from DOCX
    embedded_images = extract_embedded_images(original_file_path)

    # Add HTML content to the DOCX document
    add_html_to_docx(html_content, doc, embedded_images)

    # Save the DOCX document
    doc.save(output_file)
